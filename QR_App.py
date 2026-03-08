import io
import math
from typing import List, Tuple

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

import qrcode
from PIL import Image


EXPECTED_COLS = ["Person_Code", "Name", "Age", "Gender", "Address"]


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """
    Make column matching more forgiving (trim spaces).
    You can extend this to support aliases if needed.
    """
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    return df


def validate_columns(df: pd.DataFrame) -> Tuple[bool, List[str]]:
    missing = [c for c in EXPECTED_COLS if c not in df.columns]
    return (len(missing) == 0, missing)


def make_qr_image(data: str, box_size: int = 8, border: int = 1) -> Image.Image:
    qr = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=box_size,
        border=border,
    )
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    # Ensure it's a standard PIL Image
    return img.convert("RGB")


def pil_to_bytes_png(img: Image.Image) -> io.BytesIO:
    bio = io.BytesIO()
    img.save(bio, format="PNG")
    bio.seek(0)
    return bio


def set_a4_margins(doc: Document, top_cm=1.2, bottom_cm=1.2, left_cm=1.2, right_cm=1.2):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top_cm)
    section.bottom_margin = Cm(bottom_cm)
    section.left_margin = Cm(left_cm)
    section.right_margin = Cm(right_cm)


def set_cell_border(cell, **kwargs):
    """
    Optional: set borders (WordprocessingML). Keeps a cleaner "ID list" look.
    If you want NO borders, just don't call this.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), edge_data.get("val", "single"))
            tag.set(qn("w:sz"), str(edge_data.get("sz", 6)))
            tag.set(qn("w:space"), str(edge_data.get("space", 0)))
            tag.set(qn("w:color"), edge_data.get("color", "999999"))
            tcBorders.append(tag)


def build_docx(
    df: pd.DataFrame,
    persons_per_page: int = 8,
    qr_size_cm: float = 3.2,
    font_name: str = "Calibri",
    font_size_pt: int = 10,
    show_borders: bool = True,
) -> bytes:
    """
    persons_per_page = 6 (2x3) or 8 (2x4)
    """
    if persons_per_page not in (6, 8):
        raise ValueError("persons_per_page must be 6 or 8")

    rows_per_page = 3 if persons_per_page == 6 else 4
    cols = 2  # 2 columns per page => 6 = 2x3, 8 = 2x4

    doc = Document()
    set_a4_margins(doc)

    # Default font styling
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)

    total = len(df)
    page_count = math.ceil(total / persons_per_page)

    idx = 0
    for page in range(page_count):
        # Table for this page
        table = doc.add_table(rows=rows_per_page, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Improve layout: set each row's height by content + keep vertical center
        for r in range(rows_per_page):
            for c in range(cols):
                cell = table.cell(r, c)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # Light border (optional)
                if show_borders:
                    set_cell_border(
                        cell,
                        top={"val": "single", "sz": 6, "color": "BBBBBB"},
                        bottom={"val": "single", "sz": 6, "color": "BBBBBB"},
                        left={"val": "single", "sz": 6, "color": "BBBBBB"},
                        right={"val": "single", "sz": 6, "color": "BBBBBB"},
                    )

        # Fill cells
        for r in range(rows_per_page):
            for c in range(cols):
                if idx >= total:
                    break

                person_code = str(df.loc[idx, "Person_Code"]).strip()
                name = str(df.loc[idx, "Name"]).strip()
                age = str(df.loc[idx, "Age"]).strip()
                gender = str(df.loc[idx, "Gender"]).strip()
                address = str(df.loc[idx, "Address"]).strip()

                cell = table.cell(r, c)
                # Clear default empty paragraph
                cell.text = ""

                # Nested table: left = QR, right = text (NRC-like)
                inner = cell.add_table(rows=1, cols=2)
                inner.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Set column widths (approx)
                inner.columns[0].width = Cm(qr_size_cm + 0.3)
                inner.columns[1].width = Cm(7.5)

                # QR
                qr_img = make_qr_image(person_code, box_size=8, border=1)
                qr_bytes = pil_to_bytes_png(qr_img)

                qr_cell = inner.cell(0, 0)
                qr_cell.text = ""
                p = qr_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(qr_bytes, width=Cm(qr_size_cm))

                # Text block
                tcell = inner.cell(0, 1)
                tcell.text = ""

                def add_line(label: str, value: str, bold_label: bool = True):
                    para = tcell.add_paragraph()
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    run1 = para.add_run(f"{label}: ")
                    run1.bold = bold_label
                    para.add_run(value)

                # "NRC format" vibe: compact, label:value lines
                add_line("ID", person_code)
                add_line("Name", name)
                add_line("Age", age)
                add_line("Gender", gender)
                add_line("Address", address)

                idx += 1

        if page < page_count - 1:
            doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def read_excel(file) -> pd.DataFrame:
    # supports .xlsx/.xls
    df = pd.read_excel(file)
    df = normalize_columns(df)
    return df


# ---------------- Streamlit UI ----------------

st.set_page_config(page_title="Excel → Word IDs with QR", layout="centered")

st.title("Excel → Word ID List (QR + NRC-style text)")

st.markdown(
    """
Upload an Excel file with columns:
**Person_Code, Name, Age, Gender, Address**  
Then download a Word (.docx) file formatted in A4 with **6 or 8 persons per page**.
"""
)

uploaded = st.file_uploader("Upload Excel file (.xlsx/.xls)", type=["xlsx", "xls"])

persons_per_page = st.radio("Persons per page", options=[6, 8], horizontal=True)
qr_size_cm = st.slider("QR size (cm)", min_value=2.5, max_value=4.5, value=3.2, step=0.1)

show_borders = st.checkbox("Show light borders around each person block", value=True)

if uploaded is not None:
    try:
        df = read_excel(uploaded)

        ok, missing = validate_columns(df)
        if not ok:
            st.error(f"Missing columns: {', '.join(missing)}")
            st.stop()

        # Keep only expected columns, drop completely empty rows
        df = df[EXPECTED_COLS].copy()
        df = df.dropna(how="all")
        df = df.reset_index(drop=True)

        st.subheader("Preview")
        st.dataframe(df.head(20), use_container_width=True)

        if st.button("Generate Word file (.docx)"):
            with st.spinner("Generating..."):
                docx_bytes = build_docx(
                    df=df,
                    persons_per_page=persons_per_page,
                    qr_size_cm=qr_size_cm,
                    show_borders=show_borders,
                )

            st.success("Done!")
            st.download_button(
                label="Download Word file",
                data=docx_bytes,
                file_name="id_list.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    except Exception as e:
        st.exception(e)